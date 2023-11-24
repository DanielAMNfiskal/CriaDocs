from flask import Flask, request, jsonify, send_file
from flask_cors import CORS, cross_origin
from datetime import datetime
from python_docx_replace import docx_replace
import docx
from io import BytesIO
import zipfile

app = Flask(__name__)
CORS(app)

class DadosOBJ:
    def __init__(self, sistemas, empresa, autor, emailAutor, biblioteca):
        self.sistemas = sistemas
        self.empresa = empresa
        self.autor = autor
        self.emailAutor = emailAutor
        self.biblioteca = biblioteca

    def __str__(self) -> str:
        return f"{self.sistemas}, {self.empresa}, {self.autor}, {self.emailAutor}, {self.biblioteca}"

class Cria_Documento:
    @staticmethod
    def cria_documento_abre_site(dadosDOC):

        PDD_Abre_Site = BytesIO()
        SDD_Abre_Site = BytesIO()

        modelos_path = r"C:\Users\daniel.novaes\CriaDoc\modelosDocs"

        PDD_Abre_Site_template = docx.Document(f"{modelos_path}/PDD_MODELO_AbreSite_Fiskal_Digital.docx")
        SDD_Abre_Site_template = docx.Document(f"{modelos_path}/SDD_MODELO_AbreSite_Fiskal_Digital.docx")

        if " " in dadosDOC.sistemas:
            dadosDOC.sistemas = dadosDOC.sistemas.replace(" ", "_")

        dadosDOC.sistemas = str.upper(dadosDOC.sistemas)

        if dadosDOC.biblioteca == "1":
            dadosDOC.biblioteca = "AutoIT"
            pass
        elif dadosDOC.biblioteca == "2":
            dadosDOC.biblioteca = "Selenium"
            pass
        elif dadosDOC.biblioteca == "3":
            dadosDOC.biblioteca = "AutoIT e Selenium"
            pass
        else:
            pass

        dadosInput = {
        "empresa": str.upper(dadosDOC.empresa),
        "sistema": str.upper(dadosDOC.sistemas),
        "biblioteca": dadosDOC.biblioteca,
        "autor" : dadosDOC.autor,
        "emailAutor": dadosDOC.emailAutor,
        "data" : datetime.now().strftime("%d/%m/%Y")
        }

        docx_replace(PDD_Abre_Site_template, **dadosInput)
        docx_replace(SDD_Abre_Site_template, **dadosInput)

        PDD_Abre_Site.seek(0)
        SDD_Abre_Site.seek(0)

        try:
            PDD_Abre_Site_template.save(PDD_Abre_Site)
        except Exception as e:
            print(f"Error saving PDD_Abre_Site_template: {e}")

        try:
            SDD_Abre_Site_template.save(SDD_Abre_Site)
        except Exception as e:
            print(f"Error saving SDD_Abre_Site_template: {e}")

        return PDD_Abre_Site, SDD_Abre_Site
            
    @staticmethod
    def cria_documento_login(dadosDOC):

        PDD_Login = BytesIO()
        SDD_Login = BytesIO()

        modelos_path = r"C:\Users\daniel.novaes\CriaDoc\modelosDocs"

        PDD_Login_template = docx.Document(f"{modelos_path}/PDD_MODELO_Fiskal_Digital.docx")
        SDD_Login_template = docx.Document(f"{modelos_path}/SDD_MODELO_Fiskal_Digital.docx")

        if " " in dadosDOC.sistemas:
            dadosDOC.sistemas = dadosDOC.sistemas.replace(" ", "_")

        dadosDOC.sistemas = str.upper(dadosDOC.sistemas)

        match dadosDOC.biblioteca:
            case "1": dadosDOC.biblioteca = "AutoIT"
            case "2": dadosDOC.biblioteca = "Selenium"
            case "3": dadosDOC.biblioteca = "AutoIT e Selenium"

        dadosInput = {
        "empresa": str.upper(dadosDOC.empresa),
        "sistema": str.upper(dadosDOC.sistemas),
        "biblioteca": dadosDOC.biblioteca,
        "autor" : dadosDOC.autor,
        "emailAutor": dadosDOC.emailAutor,
        "data" : datetime.now().strftime("%d/%m/%Y")
        }

        docx_replace(PDD_Login_template, **dadosInput)
        docx_replace(SDD_Login_template, **dadosInput)

        PDD_Login.seek(0)
        SDD_Login.seek(0)

        try:
            PDD_Login_template.save(PDD_Login)
        except Exception as e:
            print(f"Error saving PDD_Abre_Site_template: {e}")

        try:
            SDD_Login_template.save(SDD_Login)
        except Exception as e:
            print(f"Error saving SDD_Abre_Site_template: {e}")

        return PDD_Login, SDD_Login


@app.route("/cria_documento/abre_site", methods=['POST'])
def cria_documento_abre_site():
    data = request.get_json()
    dados_doc = DadosOBJ(**data)
    resultado = Cria_Documento.cria_documento_abre_site(dados_doc)
    nomeDoc = str.upper(dados_doc.empresa) +"_"+ str.upper(dados_doc.sistemas)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        zip_file.writestr("PDD_Abre_Site_"+nomeDoc+".docx", resultado[0].getvalue())
        zip_file.writestr("SDD_Abre_Site_"+nomeDoc+".docx", resultado[1].getvalue())

    zip_buffer.seek(0)

    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name="documentos_{}.zip".format(nomeDoc),
        mimetype="application/zip"
    )

@app.route('/cria_documento/login', methods=['POST'])
def cria_documento_login():
    data = request.get_json()
    dados_doc = DadosOBJ(**data)
    resultado = Cria_Documento.cria_documento_login(dados_doc)
    nomeDoc = str.upper(dados_doc.empresa) +"_"+ str.upper(dados_doc.sistemas)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        zip_file.writestr("PDD_Login_"+nomeDoc+".docx", resultado[0].getvalue())
        zip_file.writestr("SDD_Login_"+nomeDoc+".docx", resultado[1].getvalue())

    zip_buffer.seek(0)

    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name="documentos_{}.zip".format(nomeDoc),
        mimetype="application/zip"
    )

if __name__ == '__main__':
    app.run()